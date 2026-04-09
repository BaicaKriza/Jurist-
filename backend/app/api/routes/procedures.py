import io
import logging
import os
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, Query, Response, UploadFile, status
from fastapi.responses import FileResponse
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.database import SessionLocal, get_db
from app.core.security import get_current_user
from app.models.user import User
from app.schemas.procedure import (
    ProcedureCreate,
    ProcedureDocumentResponse,
    ProcedureResponse,
    ProcedureSyncResponse,
    RequiredDocumentCreate,
    RequiredDocumentResponse,
    SyncRequest,
)
from app.services.procedure_service import ProcedureService

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/procedures", tags=["procedures"])

ALLOWED_DOC_TYPES = ["DST", "NJOFTIM", "SPECIFIKIME", "KRITERE", "KONTRATE", "SQARIM", "BULETIN", "TJETER"]


def _enum_value(value):
    return getattr(value, "value", value)


def _serialize_procedure(procedure, document_count: int) -> ProcedureResponse:
    return ProcedureResponse(
        id=procedure.id,
        source_name=_enum_value(procedure.source_name),
        source_url=procedure.source_url,
        reference_no=procedure.reference_no,
        notice_no=procedure.notice_no,
        authority_name=procedure.authority_name,
        object_description=procedure.object_description,
        procedure_type=procedure.procedure_type,
        contract_type=procedure.contract_type,
        cpv_code=procedure.cpv_code,
        fund_limit=float(procedure.fund_limit) if procedure.fund_limit is not None else None,
        currency=procedure.currency,
        publication_date=procedure.publication_date,
        opening_date=procedure.opening_date,
        closing_date=procedure.closing_date,
        status=_enum_value(procedure.status),
        created_at=procedure.created_at,
        updated_at=procedure.updated_at,
        document_count=document_count,
    )


def _serialize_uploaded_document(procedure_id: str, document) -> dict:
    return {
        "id": document.id,
        "title": document.title,
        "file_name": document.file_name,
        "doc_type": document.doc_type,
        "file_size": document.file_size,
        "ai_summary": document.ai_summary,
        "download_url": f"/procedures/{procedure_id}/upload/{document.id}/download",
        "created_at": document.created_at.isoformat() if document.created_at else None,
    }


def generate_ai_summary_task(doc_id: str, text: str) -> None:
    db = SessionLocal()
    try:
        import openai
        from app.models.procedure import ProcedureUploadedDocument

        client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": (
                        "Analizoni kete dokument prokurimi shqiptar. Nxirrni dhe permbledhni: "
                        "lloji i dokumentit, kerkesat kryesore, kriteret e kualifikimit, afatet, "
                        "kushtet speciale. Ktheni si JSON: "
                        '{"doc_type":"...","summary":"...","key_points":[],"deadlines":[],"requirements":[],"flags":[]} '
                        f"Teksti: {text[:4000]}"
                    ),
                }
            ],
            response_format={"type": "json_object"},
        )
        summary = response.choices[0].message.content
        document = db.execute(
            select(ProcedureUploadedDocument).where(ProcedureUploadedDocument.id == doc_id)
        ).scalar_one_or_none()
        if document:
            document.ai_summary = summary
            db.commit()
            logger.info("AI summary generated for procedure document %s", doc_id)
    except Exception as exc:
        logger.warning("AI summary generation failed for doc %s: %s", doc_id, exc)
    finally:
        db.close()


@router.get(
    "",
    response_model=dict,
    summary="List procurement procedures",
)
def list_procedures(
    source_name: Optional[str] = Query(None),
    status_filter: Optional[str] = Query(None, alias="status"),
    authority_name: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> dict:
    service = ProcedureService(db)
    procedures, total = service.get_procedures(
        source_name=source_name,
        status_filter=status_filter,
        authority_name=authority_name,
        search=search,
        page=page,
        page_size=page_size,
    )
    items = [_serialize_procedure(proc, service.get_procedure_doc_count(proc.id)) for proc in procedures]
    return {
        "items": items,
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": (total + page_size - 1) // page_size,
    }


@router.get(
    "/{procedure_id}",
    response_model=ProcedureResponse,
    summary="Get a procedure by ID",
)
def get_procedure(
    procedure_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> ProcedureResponse:
    service = ProcedureService(db)
    procedure = service.get_procedure(procedure_id)
    return _serialize_procedure(procedure, service.get_procedure_doc_count(procedure_id))


@router.post("/sync", response_model=ProcedureSyncResponse, summary="Sync from APP portal")
async def sync_procedures(
    payload: SyncRequest,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> ProcedureSyncResponse:
    service = ProcedureService(db)
    return await service.sync_from_app(payload)


@router.post("/{procedure_id}/analyze", response_model=dict, summary="Analyze with AI")
async def analyze_procedure(
    procedure_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> dict:
    service = ProcedureService(db)
    analysis = await service.analyze_procedure(procedure_id)
    return {
        "id": analysis.id,
        "procedure_id": analysis.procedure_id,
        "analysis_type": analysis.analysis_type,
        "summary": analysis.summary,
        "legal_notes": analysis.legal_notes,
        "technical_notes": analysis.technical_notes,
        "financial_notes": analysis.financial_notes,
        "risk_level": _enum_value(analysis.risk_level),
        "recommended_action": analysis.recommended_action,
        "created_at": analysis.created_at.isoformat() if analysis.created_at else None,
    }


@router.post("/{procedure_id}/download-documents", response_model=dict)
async def download_procedure_documents(
    procedure_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> dict:
    service = ProcedureService(db)
    return await service.download_procedure_documents(procedure_id)


@router.get(
    "/{procedure_id}/documents",
    response_model=list[ProcedureDocumentResponse],
    summary="Get scraped documents for a procedure",
)
def get_procedure_documents(
    procedure_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> list[ProcedureDocumentResponse]:
    service = ProcedureService(db)
    docs = service.get_procedure_documents(procedure_id)
    return [ProcedureDocumentResponse.model_validate(doc) for doc in docs]


@router.post(
    "/{procedure_id}/upload",
    summary="Upload a file directly to a procedure",
    status_code=status.HTTP_201_CREATED,
)
async def upload_procedure_document(
    procedure_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str = Form(...),
    doc_type: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from app.models.procedure import Procedure as ProcedureModel
    from app.models.procedure import ProcedureUploadedDocument

    if doc_type not in ALLOWED_DOC_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"doc_type i pavlefshem. Vlerat e lejuara: {', '.join(ALLOWED_DOC_TYPES)}",
        )

    procedure = db.execute(
        select(ProcedureModel).where(ProcedureModel.id == procedure_id)
    ).scalar_one_or_none()
    if not procedure:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Procedura nuk u gjet")

    content = await file.read()
    file_size = len(content)

    storage_root = os.getenv("STORAGE_PATH") or "./storage"
    upload_dir = storage_root if storage_root.endswith("procedure_docs") else os.path.join(storage_root, "procedure_docs")
    os.makedirs(upload_dir, exist_ok=True)

    original_name = os.path.basename(file.filename or "document")
    safe_name = f"{procedure_id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{original_name}"
    file_path = os.path.join(upload_dir, safe_name)
    with open(file_path, "wb") as output_file:
        output_file.write(content)

    extracted_text = ""
    mime_type = file.content_type or ""
    try:
        if "pdf" in mime_type or original_name.lower().endswith(".pdf"):
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                extracted_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif "wordprocessingml" in mime_type or original_name.lower().endswith(".docx"):
            from docx import Document as DocxDocument

            document = DocxDocument(io.BytesIO(content))
            extracted_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    except Exception as exc:
        logger.warning("Text extraction failed: %s", exc)

    uploaded_doc = ProcedureUploadedDocument(
        procedure_id=procedure_id,
        file_name=original_name,
        file_path=file_path,
        doc_type=doc_type,
        title=title,
        extracted_text=extracted_text[:50000] if extracted_text else None,
        file_size=file_size,
        mime_type=mime_type,
        is_deleted=False,
        created_at=datetime.utcnow(),
    )
    db.add(uploaded_doc)
    db.commit()
    db.refresh(uploaded_doc)

    if extracted_text and os.getenv("OPENAI_API_KEY"):
        background_tasks.add_task(generate_ai_summary_task, uploaded_doc.id, extracted_text)

    logger.info("Procedure document uploaded by user %s: %s", current_user.id, uploaded_doc.id)
    return {
        "id": uploaded_doc.id,
        "file_name": uploaded_doc.file_name,
        "doc_type": uploaded_doc.doc_type,
        "title": uploaded_doc.title,
        "download_url": f"/procedures/{procedure_id}/upload/{uploaded_doc.id}/download",
        "file_size": uploaded_doc.file_size,
        "ai_summary": None,
        "created_at": uploaded_doc.created_at.isoformat() if uploaded_doc.created_at else None,
    }


@router.get("/{procedure_id}/upload", summary="List uploaded documents for a procedure")
def list_uploaded_documents(
    procedure_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    from app.models.procedure import ProcedureUploadedDocument

    documents = db.execute(
        select(ProcedureUploadedDocument)
        .where(
            ProcedureUploadedDocument.procedure_id == procedure_id,
            ProcedureUploadedDocument.is_deleted.is_(False),
        )
        .order_by(ProcedureUploadedDocument.created_at.desc())
    ).scalars().all()

    return [_serialize_uploaded_document(procedure_id, document) for document in documents]


@router.delete(
    "/{procedure_id}/upload/{doc_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Soft-delete an uploaded procedure document",
)
def delete_uploaded_document(
    procedure_id: str,
    doc_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    from app.models.procedure import ProcedureUploadedDocument

    document = db.execute(
        select(ProcedureUploadedDocument).where(
            ProcedureUploadedDocument.id == doc_id,
            ProcedureUploadedDocument.procedure_id == procedure_id,
        )
    ).scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Dokumenti nuk u gjet")

    document.is_deleted = True
    db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.get(
    "/{procedure_id}/upload/{doc_id}/download",
    summary="Download an uploaded procedure document",
)
def download_uploaded_document(
    procedure_id: str,
    doc_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    from app.models.procedure import ProcedureUploadedDocument

    document = db.execute(
        select(ProcedureUploadedDocument).where(
            ProcedureUploadedDocument.id == doc_id,
            ProcedureUploadedDocument.procedure_id == procedure_id,
            ProcedureUploadedDocument.is_deleted.is_(False),
        )
    ).scalar_one_or_none()
    if not document or not os.path.exists(document.file_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Skedari nuk u gjet")

    return FileResponse(
        document.file_path,
        filename=document.file_name,
        media_type=document.mime_type or "application/octet-stream",
    )


@router.post(
    "",
    response_model=ProcedureResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Manually create a procedure",
)
def create_procedure(
    payload: ProcedureCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> ProcedureResponse:
    from app.models.procedure import Procedure as ProcedureModel

    if payload.reference_no:
        existing = db.execute(
            select(ProcedureModel).where(ProcedureModel.reference_no == payload.reference_no)
        ).scalar_one_or_none()
        if existing:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"Procedura me referencen '{payload.reference_no}' ekziston tashme",
            )

    procedure = ProcedureModel(
        source_name=payload.source_name,
        source_url=payload.source_url,
        reference_no=payload.reference_no,
        notice_no=payload.notice_no,
        authority_name=payload.authority_name,
        object_description=payload.object_description,
        procedure_type=payload.procedure_type,
        contract_type=payload.contract_type,
        cpv_code=payload.cpv_code,
        fund_limit=payload.fund_limit,
        currency=payload.currency or "ALL",
        publication_date=payload.publication_date,
        opening_date=payload.opening_date,
        closing_date=payload.closing_date,
        status=payload.status,
    )
    db.add(procedure)
    db.commit()
    db.refresh(procedure)
    logger.info("Procedure '%s' created by user %s", procedure.reference_no, current_user.id)
    return _serialize_procedure(procedure, 0)


@router.post(
    "/{procedure_id}/requirements",
    response_model=RequiredDocumentResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Add a required document to a procedure",
)
def add_requirement(
    procedure_id: str,
    payload: RequiredDocumentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> RequiredDocumentResponse:
    from app.models.analysis import RequiredDocumentItem
    from app.models.procedure import Procedure as ProcedureModel

    procedure = db.execute(
        select(ProcedureModel).where(ProcedureModel.id == procedure_id)
    ).scalar_one_or_none()
    if not procedure:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Procedura nuk u gjet")

    item = RequiredDocumentItem(
        procedure_id=procedure_id,
        name=payload.name,
        category=payload.category,
        description=payload.description,
        mandatory=payload.resolved_mandatory,
        issuer_type=payload.issuer_type,
        source_hint=payload.source_hint,
        validity_rule=payload.validity_rule,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    logger.info("Requirement '%s' added to procedure %s by %s", item.name, procedure_id, current_user.id)
    return RequiredDocumentResponse.model_validate(item)


@router.get(
    "/{procedure_id}/requirements",
    response_model=list[RequiredDocumentResponse],
    summary="List required documents for a procedure",
)
def list_requirements(
    procedure_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> list[RequiredDocumentResponse]:
    from app.models.analysis import RequiredDocumentItem

    items = db.execute(
        select(RequiredDocumentItem)
        .where(RequiredDocumentItem.procedure_id == procedure_id)
        .order_by(RequiredDocumentItem.category, RequiredDocumentItem.name)
    ).scalars().all()
    return [RequiredDocumentResponse.model_validate(item) for item in items]


@router.delete(
    "/{procedure_id}/requirements/{requirement_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete a required document from a procedure",
)
def delete_requirement(
    procedure_id: str,
    requirement_id: str,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> None:
    from app.models.analysis import RequiredDocumentItem

    item = db.execute(
        select(RequiredDocumentItem).where(
            RequiredDocumentItem.id == requirement_id,
            RequiredDocumentItem.procedure_id == procedure_id,
        )
    ).scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Kerkesa nuk u gjet")

    db.delete(item)
    db.commit()
    logger.info("Requirement %s deleted from procedure %s", requirement_id, procedure_id)


@router.post(
    "/{procedure_id}/documents/upload",
    status_code=status.HTTP_201_CREATED,
    include_in_schema=False,
    summary="Alias: upload doc to procedure",
)
async def upload_procedure_document_alias(
    procedure_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str = Form(...),
    doc_type: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return await upload_procedure_document(
        procedure_id=procedure_id,
        background_tasks=background_tasks,
        file=file,
        title=title,
        doc_type=doc_type,
        db=db,
        current_user=current_user,
    )


@router.get(
    "/{procedure_id}/documents/list",
    include_in_schema=False,
    summary="Alias: list uploaded docs",
)
def list_uploaded_documents_alias(
    procedure_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return list_uploaded_documents(procedure_id=procedure_id, db=db, _=current_user)
